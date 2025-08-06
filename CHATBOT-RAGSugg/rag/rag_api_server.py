# File: rag_api_server.py
from fastapi import FastAPI, UploadFile, Form , File
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from PyPDF2 import PdfReader
import os
import uuid
from docx import Document
from fastapi.responses import FileResponse
from fastapi import HTTPException



from setup_and_run import run_rag_on_text

app = FastAPI()
EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)
# Allow CORS for BC/JS/ControlAddIn access
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # You can restrict this later
    allow_methods=["*"],
    allow_headers=["*"],
)

from fastapi import File

@app.post("/copilot/rag-process")
async def rag_process(prompt: str = Form(...), file: UploadFile = File(None)):
    try:
        context = ""
        if file and file.filename.endswith(".pdf"):
            reader = PdfReader(file.file)
            context = "\n\n".join([page.extract_text() or "" for page in reader.pages])
            file.file.seek(0)
        elif file and file.filename.endswith(".txt"):
            contents = await file.read()
            context = contents.decode("utf-8")

        # Call your existing RAG pipeline
        result_text = run_rag_on_text(prompt, context)

        # Save the result as a docx file
        docx_filename = save_response_as_docx(result_text)

        return {
            "success": True,
            "response_text": result_text,
            "filename": docx_filename  # Return the file name for download link later
        }

    except Exception as e:
        return JSONResponse(status_code=500, content={
            "success": False,
            "error": str(e)
        })
    
@app.get("/copilot/download-docx/{filename}")
async def download_docx(filename: str):
    filepath = os.path.join(EXPORT_DIR, filename)

    if not os.path.isfile(filepath):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=filepath,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

def save_response_as_docx(text: str, export_dir: str = EXPORT_DIR) -> str:
    filename = f"rag_output_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(export_dir, filename)

    doc = Document()
    doc.add_heading("Generated AI Response", level=1)
    doc.add_paragraph(text)
    doc.save(filepath)

    return filename
